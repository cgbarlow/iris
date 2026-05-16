"""v6.2.0 (ADR-179, SPEC-179-A): md → docx renderer tests.

Exercises `app/export/renderers/docx.py` end-to-end by rendering a
markdown fixture and re-parsing the produced bytes with python-docx
to verify structure made it through.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document  # type: ignore[import-untyped]

from app.export.renderers import docx as docx_renderer


def _render_and_reopen(markdown: str, title: str = "Test"):
    data, filename = docx_renderer.render(markdown, title)
    assert filename.endswith(".docx")
    assert data.startswith(b"PK\x03\x04"), "docx must be a ZIP (PK header)"
    return Document(BytesIO(data))


def test_title_appears_as_top_heading() -> None:
    doc = _render_and_reopen("Body text.", title="My Title")
    # First paragraph should be the title heading.
    first = doc.paragraphs[0]
    assert "My Title" in first.text


def test_simple_paragraph_round_trips() -> None:
    doc = _render_and_reopen("Hello world")
    texts = [p.text for p in doc.paragraphs]
    assert any("Hello world" in t for t in texts)


def test_heading_h1_h2_h3() -> None:
    md = "# H1\n\n## H2\n\n### H3\n"
    doc = _render_and_reopen(md, title="Headings")
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles


def test_bullet_list_round_trips() -> None:
    md = "- one\n- two\n- three\n"
    doc = _render_and_reopen(md, title="Bullets")
    bullet_items = [
        p.text for p in doc.paragraphs
        if "List Bullet" in p.style.name
    ]
    assert "one" in bullet_items
    assert "two" in bullet_items
    assert "three" in bullet_items


def test_ordered_list_round_trips() -> None:
    md = "1. first\n2. second\n3. third\n"
    doc = _render_and_reopen(md, title="Numbers")
    number_items = [
        p.text for p in doc.paragraphs
        if "List Number" in p.style.name
    ]
    assert "first" in number_items
    assert "second" in number_items
    assert "third" in number_items


def test_code_block_preserved() -> None:
    md = "```\nlet x = 42\nprintln(x)\n```\n"
    doc = _render_and_reopen(md, title="Code")
    texts = [p.text for p in doc.paragraphs]
    joined = " ".join(texts)
    assert "let x = 42" in joined
    assert "println(x)" in joined


def test_mermaid_block_passes_through_verbatim() -> None:
    md = (
        "Some prose.\n\n"
        "```mermaid\n"
        "flowchart LR\n"
        "  A --> B\n"
        "```\n"
    )
    doc = _render_and_reopen(md, title="Mermaid")
    joined = " ".join(p.text for p in doc.paragraphs)
    assert "flowchart LR" in joined
    assert "A --> B" in joined


def test_empty_markdown_still_produces_valid_docx() -> None:
    doc = _render_and_reopen("", title="Empty")
    # Just the title heading.
    assert any("Empty" in p.text for p in doc.paragraphs)


def test_filename_slug_derives_from_title() -> None:
    _, filename = docx_renderer.render("body", title="Banana Monoculture")
    assert filename.startswith("banana-monoculture-")
    assert filename.endswith(".docx")
