"""Markdown → DOCX renderer (ADR-179, v6.2.0).

Uses `markdown-it-py` to tokenise the markdown source, then walks the
token stream emitting `python-docx` paragraphs / headings / lists /
code blocks. Mermaid fenced blocks pass through verbatim as
monospace code blocks — the docx reader can paste them into a
mermaid-rendering tool.

Recipe modelled on Anthropic's `skills/docx` (Apache 2.0) — minimal
adapter rather than a vendored copy.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt
from markdown_it import MarkdownIt

from app.export.renderers._common import slug_filename


_HEADING_STYLE = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 5",
    "h6": "Heading 6",
}


def render(markdown: str, title: str) -> tuple[bytes, str]:
    """Render markdown to a docx bytes blob.

    Always starts with the document title as a Title-styled paragraph,
    followed by the markdown content rendered structurally.
    """
    doc = Document()
    doc.add_heading(title or "Untitled", level=0)

    md = MarkdownIt("commonmark", {"breaks": False, "html": False})
    tokens = md.parse(markdown or "")

    _walk_tokens(doc, tokens)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), slug_filename(title, "docx")


def _walk_tokens(doc: Document, tokens: list) -> None:  # noqa: C901, PLR0912, PLR0915
    """Emit docx paragraphs / lists / code blocks from the token stream.

    The walker maintains a simple stack for list nesting and a
    'current paragraph' pointer for inline content accumulation.
    Inline images / links are flattened to their alt text / URL — the
    docx renderer does not embed binary images (keeps the renderer
    container small; mermaid stays as text).
    """
    list_stack: list[str] = []  # 'bullet' | 'number' nesting
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t.type == "heading_open":
            style = _HEADING_STYLE.get(t.tag, "Heading 1")
            text = _collect_inline_text(tokens[i + 1])
            doc.add_paragraph(text, style=style)
            i += 3  # heading_open + inline + heading_close
            continue

        if t.type == "paragraph_open":
            inline = tokens[i + 1]
            text = _collect_inline_text(inline)
            if list_stack:
                style = (
                    "List Bullet" if list_stack[-1] == "bullet" else "List Number"
                )
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(text)
            i += 3
            continue

        if t.type == "fence" or t.type == "code_block":
            code = t.content.rstrip("\n")
            for line in code.split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            i += 1
            continue

        if t.type == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue
        if t.type == "ordered_list_open":
            list_stack.append("number")
            i += 1
            continue
        if t.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if t.type == "blockquote_open":
            # Find the matching close, render contents indented.
            depth = 1
            j = i + 1
            inner: list[str] = []
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                elif tokens[j].type == "inline":
                    inner.append(tokens[j].content)
                j += 1
            doc.add_paragraph("> " + " ".join(inner), style="Intense Quote")
            i = j
            continue

        if t.type == "hr":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Unrecognised top-level token — advance.
        i += 1


def _collect_inline_text(inline_token) -> str:
    """Flatten an inline token into plain text.

    Code spans wrapped in backticks, link URLs surfaced inline, alt
    text used for images. Bold/italic markers dropped — docx run
    formatting is not threaded through here to keep the walker
    simple.
    """
    if inline_token is None or inline_token.type != "inline":
        return ""
    parts: list[str] = []
    for child in (inline_token.children or []):
        if child.type == "text":
            parts.append(child.content)
        elif child.type == "code_inline":
            parts.append(f"`{child.content}`")
        elif child.type == "link_open":
            href = next(
                (a[1] for a in (child.attrs or []) if a[0] == "href"),
                "",
            ) if isinstance(child.attrs, list) else child.attrs.get("href", "")
            parts.append(f"[")
            # Link text is in subsequent children until link_close.
            # The simpler approach: accumulate inside _collect_inline_text
            # by tracking a link buffer. For now, surface the URL inline.
            parts.append(f"]({href})")
        elif child.type == "softbreak" or child.type == "hardbreak":
            parts.append(" ")
        elif child.type == "image":
            alt = child.content or ""
            src = next(
                (a[1] for a in (child.attrs or []) if a[0] == "src"),
                "",
            ) if isinstance(child.attrs, list) else child.attrs.get("src", "")
            parts.append(f"[image: {alt or src}]")
    return "".join(parts).strip()
